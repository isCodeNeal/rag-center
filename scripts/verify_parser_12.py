"""Verify multi-format upload and retrieval pipeline.

Usage:
  python scripts/verify_parser_12.py --api-key <key> --kb-id <kb_id>
  python scripts/verify_parser_12.py --base-url http://localhost:8000 --api-key <key> --kb-id <kb_id>

Tests:
  1. JSON upload (.md) -> poll SUCCESS -> retrieve keyword -> PASS/FAIL
  2. Multipart upload (.docx) -> poll SUCCESS -> retrieve keyword -> PASS/FAIL (skipped if python-docx missing)
  3. Upload unsupported extension (.zip) -> expect PARAM_ERROR -> PASS/FAIL
"""
from __future__ import annotations

import argparse
import os
import sys
import tempfile
import time
from pathlib import Path
from typing import Optional

try:
    import requests
except ImportError:
    sys.exit("requests is required: pip install requests")

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Document status constants (matches backend enums)
STATUS_SUCCESS = 1
STATUS_FAILED = 2
STATUS_PROCESSING = 3

POLL_INTERVAL = 3   # seconds between status polls
POLL_TIMEOUT = 60   # max seconds to wait for indexing


def make_session(api_key: Optional[str]) -> requests.Session:
    s = requests.Session()
    if api_key:
        s.headers["Authorization"] = f"Bearer {api_key}"
    return s


def poll_document_status(
    session: requests.Session,
    base_url: str,
    document_id: str,
) -> int:
    """Poll document status until SUCCESS/FAILED or timeout. Returns final status code."""
    deadline = time.monotonic() + POLL_TIMEOUT
    while time.monotonic() < deadline:
        resp = session.get(f"{base_url}/api/v1/documents/{document_id}")
        resp.raise_for_status()
        body = resp.json()
        status = body.get("data", {}).get("status", STATUS_PROCESSING)
        if status in (STATUS_SUCCESS, STATUS_FAILED):
            return status
        time.sleep(POLL_INTERVAL)
    return STATUS_PROCESSING  # timeout


def retrieve_keyword(
    session: requests.Session,
    base_url: str,
    kb_id: str,
    keyword: str,
) -> list[dict]:
    """Run a retrieval query and return chunks list."""
    payload = {
        "kb_id": kb_id,
        "user_id": "verify_script",
        "query": keyword,
        "top_k": 3,
    }
    resp = session.post(f"{base_url}/api/v1/rag/retrieve", json=payload)
    resp.raise_for_status()
    body = resp.json()
    return body.get("data", {}).get("retrieved_chunks", [])


def upload_json(
    session: requests.Session,
    base_url: str,
    kb_id: str,
    title: str,
    content: str,
) -> dict:
    """Upload a text document via JSON body."""
    resp = session.post(
        f"{base_url}/api/v1/documents/upload",
        json={"kb_id": kb_id, "title": title, "content": content},
    )
    return resp


def upload_multipart(
    session: requests.Session,
    base_url: str,
    kb_id: str,
    title: str,
    file_path: Path,
) -> requests.Response:
    """Upload a binary document via multipart/form-data."""
    with open(file_path, "rb") as fh:
        resp = session.post(
            f"{base_url}/api/v1/documents/upload",
            files={"file": (file_path.name, fh)},
            data={"kb_id": kb_id, "title": title},
        )
    return resp


def run_tests(base_url: str, api_key: Optional[str], kb_id: str) -> None:
    session = make_session(api_key)
    passed = 0
    failed = 0

    with tempfile.TemporaryDirectory(prefix="verify_parser_12_") as tmpdir:
        tmp = Path(tmpdir)

        # ── Test 1: JSON upload of .md ─────────────────────────────────────
        print("\n[1] JSON upload (.md) ...")
        md_keyword = "量子纠缠验收关键词"
        md_content = (
            "# 验收测试文档\n\n"
            "## 量子纠缠简介\n\n"
            f"本文档用于验证 RAG 管道的 md 上传路径。关键词：{md_keyword}。\n"
        )
        md_file = tmp / "sample.md"
        md_file.write_text(md_content, encoding="utf-8")

        try:
            resp = upload_json(session, base_url, kb_id, "verify_md", md_content)
            body = resp.json()
            if body.get("code") != 0:
                raise RuntimeError(f"upload failed: {body.get('msg')}")
            doc_id = body["data"]["document_id"]
            print(f"   document_id={doc_id}  polling status ...")
            status = poll_document_status(session, base_url, doc_id)
            if status != STATUS_SUCCESS:
                raise RuntimeError(f"indexing did not reach SUCCESS (status={status})")
            chunks = retrieve_keyword(session, base_url, kb_id, md_keyword)
            if not chunks:
                raise RuntimeError("retrieve returned 0 chunks for md keyword")
            print(f"   retrieved {len(chunks)} chunk(s)  -> PASS")
            passed += 1
        except Exception as exc:
            print(f"   ERROR: {exc}  -> FAIL")
            failed += 1

        # ── Test 2: Multipart upload of .docx ─────────────────────────────
        if not HAS_DOCX:
            print("\n[2] Multipart upload (.docx) ... SKIPPED (python-docx not installed)")
        else:
            print("\n[2] Multipart upload (.docx) ...")
            docx_keyword = "DocxVerifyUniqueToken9527"
            docx_file = tmp / "sample.docx"
            doc = DocxDocument()
            doc.add_heading("验收测试 DOCX", level=1)
            doc.add_paragraph(f"本段落包含关键词：{docx_keyword}，用于验证 docx 解析路径。")
            doc.save(str(docx_file))

            try:
                resp = upload_multipart(session, base_url, kb_id, "verify_docx", docx_file)
                body = resp.json()
                if body.get("code") != 0:
                    raise RuntimeError(f"upload failed: {body.get('msg')}")
                doc_id = body["data"]["document_id"]
                print(f"   document_id={doc_id}  polling status ...")
                status = poll_document_status(session, base_url, doc_id)
                if status != STATUS_SUCCESS:
                    raise RuntimeError(f"indexing did not reach SUCCESS (status={status})")
                chunks = retrieve_keyword(session, base_url, kb_id, docx_keyword)
                if not chunks:
                    raise RuntimeError("retrieve returned 0 chunks for docx keyword")
                print(f"   retrieved {len(chunks)} chunk(s)  -> PASS")
                passed += 1
            except Exception as exc:
                print(f"   ERROR: {exc}  -> FAIL")
                failed += 1

        # ── Test 3: Unsupported extension (.zip) → expect PARAM_ERROR ─────
        print("\n[3] Upload unsupported extension (.zip) ...")
        zip_file = tmp / "bad.zip"
        zip_file.write_bytes(b"PK\x03\x04")  # minimal ZIP magic bytes

        try:
            resp = upload_multipart(session, base_url, kb_id, "verify_zip", zip_file)
            body = resp.json()
            code = body.get("code", 0)
            # Expect a non-zero error code (PARAM_ERROR or similar rejection)
            if code != 0:
                print(f"   server rejected .zip with code={code} msg={body.get('msg')}  -> PASS")
                passed += 1
            else:
                raise RuntimeError(
                    f"expected error for .zip but got code=0, data={body.get('data')}"
                )
        except Exception as exc:
            print(f"   ERROR: {exc}  -> FAIL")
            failed += 1

    # ── Summary ────────────────────────────────────────────────────────────
    total = passed + failed
    print(f"\n{'='*40}")
    print(f"Results: {passed}/{total} passed, {failed} failed")
    if failed:
        print("OVERALL: FAIL")
        sys.exit(1)
    else:
        print("OVERALL: PASS")
        sys.exit(0)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Verify multi-format upload and retrieval pipeline"
    )
    parser.add_argument(
        "--base-url",
        default="http://localhost:8000",
        help="API base URL (default: http://localhost:8000)",
    )
    parser.add_argument(
        "--api-key",
        default=None,
        help="Bearer API key (Authorization header)",
    )
    parser.add_argument(
        "--kb-id",
        required=True,
        help="Knowledge base ID to use for test uploads",
    )
    args = parser.parse_args()
    run_tests(base_url=args.base_url.rstrip("/"), api_key=args.api_key, kb_id=args.kb_id)
